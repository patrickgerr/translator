from deep_translator import (GoogleTranslator)
from docx import Document

doc = Document('me.docx')
f= doc.paragraphs
for p in doc.paragraphs:
    text=p.text
    translated = GoogleTranslator(source='en', target='ta').translate(text=text)
    print(translated)
    f = open("demo.txt", "a")
    f.write(translated)
    f.close()